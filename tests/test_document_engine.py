"""
æ–‡ä»¶è™•ç†å¼•æ“ & é€²éšæª¢ç´¢ â€” å–®å…ƒæ¸¬è©¦å¥—ä»¶

æ¸¬è©¦ç¯„åœï¼š
  1. DocumentParser â€” å¤šæ ¼å¼è§£æ + å“è³ªå ±å‘Š
  2. TextChunker â€” ç²¾ç¢º Token åˆ‡ç‰‡ + ç« ç¯€åµæ¸¬
  3. KnowledgeBaseRetriever â€” é€²éšæª¢ç´¢åŠŸèƒ½é©—è­‰
"""

import os
import sys
import json
import tempfile
import time

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Colored output helpers
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
CYAN = "\033[96m"
BOLD = "\033[1m"
RESET = "\033[0m"

passed = 0
failed = 0
warnings = 0


def test(name: str, condition: bool, detail: str = ""):
    global passed, failed
    if condition:
        passed += 1
        print(f"  {GREEN}âœ“{RESET} {name}")
    else:
        failed += 1
        print(f"  {RED}âœ—{RESET} {name}")
        if detail:
            print(f"    {RED}â†’ {detail}{RESET}")


def warn(msg: str):
    global warnings
    warnings += 1
    print(f"  {YELLOW}âš {RESET} {msg}")


def section(title: str):
    print(f"\n{BOLD}{CYAN}{'=' * 60}{RESET}")
    print(f"{BOLD}{CYAN}{title}{RESET}")
    print(f"{BOLD}{CYAN}{'=' * 60}{RESET}\n")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# 1. DocumentParser æ¸¬è©¦
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

section("1. DocumentParser â€” æ ¼å¼æ”¯æ´èˆ‡åµæ¸¬")

from app.services.document_parser import (
    DocumentParser,
    TextChunker,
    QualityReport,
    SUPPORTED_FORMATS,
    _HAS_TIKTOKEN,
    _HAS_PDFPLUMBER,
    _HAS_OPENPYXL,
    _HAS_OCR,
    _HAS_CHARDET,
    _HAS_RTF,
)

# 1.1 æ ¼å¼æ˜ å°„å®Œæ•´æ€§
test("SUPPORTED_FORMATS åŒ…å« PDF",    ".pdf" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« DOCX",   ".docx" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« DOC",    ".doc" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« TXT",    ".txt" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« XLSX",   ".xlsx" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« XLS",    ".xls" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« CSV",    ".csv" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« HTML",   ".html" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« HTM",    ".htm" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« MD",     ".md" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« RTF",    ".rtf" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« JSON",   ".json" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« JPG",    ".jpg" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« PNG",    ".png" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« TIFF",   ".tiff" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS åŒ…å« BMP",    ".bmp" in SUPPORTED_FORMATS)
test("æ”¯æ´æ ¼å¼ç¸½æ•¸ >= 16",            len(SUPPORTED_FORMATS) >= 16)

# 1.2 detect_file_type
test("detect_file_type('.pdf') = 'pdf'",    DocumentParser.detect_file_type("test.pdf") == "pdf")
test("detect_file_type('.xlsx') = 'xlsx'",  DocumentParser.detect_file_type("data.xlsx") == "xlsx")
test("detect_file_type('.html') = 'html'",  DocumentParser.detect_file_type("page.html") == "html")
test("detect_file_type('.csv') = 'csv'",    DocumentParser.detect_file_type("data.csv") == "csv")
test("detect_file_type('.md') = 'markdown'", DocumentParser.detect_file_type("README.md") == "markdown")
test("detect_file_type('.json') = 'json'",  DocumentParser.detect_file_type("config.json") == "json")
test("detect_file_type('.jpg') = 'image'",  DocumentParser.detect_file_type("photo.jpg") == "image")

try:
    DocumentParser.detect_file_type("test.xyz")
    test("ä¸æ”¯æ´æ ¼å¼æ‹‹ ValueError", False)
except ValueError:
    test("ä¸æ”¯æ´æ ¼å¼æ‹‹ ValueError", True)

# 1.3 ä¾è³´åµæ¸¬
print(f"\n  {CYAN}ä¾è³´åµæ¸¬ç‹€æ…‹ï¼š{RESET}")
print(f"    tiktoken:     {'âœ“ å·²å®‰è£' if _HAS_TIKTOKEN else 'âœ— æœªå®‰è£'}")
print(f"    pdfplumber:   {'âœ“ å·²å®‰è£' if _HAS_PDFPLUMBER else 'âœ— æœªå®‰è£'}")
print(f"    openpyxl:     {'âœ“ å·²å®‰è£' if _HAS_OPENPYXL else 'âœ— æœªå®‰è£'}")
print(f"    chardet:      {'âœ“ å·²å®‰è£' if _HAS_CHARDET else 'âœ— æœªå®‰è£'}")
print(f"    striprtf:     {'âœ“ å·²å®‰è£' if _HAS_RTF else 'âœ— æœªå®‰è£'}")
print(f"    pytesseract:  {'âœ“ å·²å®‰è£' if _HAS_OCR else 'âœ— æœªå®‰è£'}")

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# 2. å¯¦éš›æ–‡ä»¶è§£ææ¸¬è©¦ï¼ˆä½¿ç”¨è‡¨æ™‚æ–‡ä»¶ï¼‰
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

section("2. DocumentParser â€” å¯¦éš›è§£ææ¸¬è©¦")

# 2.1 TXT è§£æ
with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False) as f:
    f.write("é€™æ˜¯ä¸€æ®µæ¸¬è©¦æ–‡å­—ã€‚\n\nå“¡å·¥è«‹å‡è¾¦æ³•å¦‚ä¸‹ï¼š\n\nç¬¬ä¸€æ¢ã€é©ç”¨ç¯„åœï¼šæœ¬è¾¦æ³•é©ç”¨æ–¼å…¨é«”æ­£å¼å“¡å·¥ã€‚\nç¬¬äºŒæ¢ã€è«‹å‡ç¨®é¡ï¼šç—…å‡ã€äº‹å‡ã€ç‰¹ä¼‘å‡ã€å©šå‡ã€å–ªå‡ã€ç”¢å‡ã€‚\nç¬¬ä¸‰æ¢ã€ç—…å‡è¦å®šï¼šæ¯å¹´ç—…å‡ä¸è¶…éä¸‰åæ—¥ï¼Œä»¥åŠè–ªè¨ˆç®—ã€‚")
    txt_path = f.name

text, meta = DocumentParser.parse(txt_path, "txt")
test("TXT è§£æ: æ–‡å­—å…§å®¹é•·åº¦ > 50",     len(text) > 50)
test("TXT è§£æ: quality_level ä¸ç‚º failed", meta.get("quality_level") != "failed")
test("TXT è§£æ: åŒ…å« 'è«‹å‡'",            "è«‹å‡" in text)
test("TXT è§£æ: parse_time_ms å­˜åœ¨",      "parse_time_ms" in meta)
os.unlink(txt_path)

# 2.2 Markdown è§£æ
with tempfile.NamedTemporaryFile(suffix=".md", mode="w", encoding="utf-8", delete=False) as f:
    f.write("# å“¡å·¥æ‰‹å†Š\n\n## ç¬¬ä¸€ç«  ç¸½å‰‡\n\næœ¬æ‰‹å†Šé©ç”¨æ–¼å…¨é«”å“¡å·¥ã€‚\n\n## ç¬¬äºŒç«  å‡ºå‹¤\n\nä¸Šåˆä¹é»ä¸Šç­ï¼Œä¸‹åˆå…­é»ä¸‹ç­ã€‚\n\n### 2.1 å½ˆæ€§å·¥æ™‚\n\nå¯ç”³è«‹å½ˆæ€§å·¥æ™‚ã€‚")
    md_path = f.name

text, meta = DocumentParser.parse(md_path, "markdown")
test("MD è§£æ: åŒ…å«æ¨™é¡Œ 'å“¡å·¥æ‰‹å†Š'",     "å“¡å·¥æ‰‹å†Š" in text)
test("MD è§£æ: format_detected = 'markdown'", meta.get("format_detected") == "markdown")
os.unlink(md_path)

# 2.3 CSV è§£æ
with tempfile.NamedTemporaryFile(suffix=".csv", mode="w", encoding="utf-8", delete=False) as f:
    f.write("å§“å,éƒ¨é–€,å‡åˆ¥,å¤©æ•¸\nå¼µä¸‰,å·¥ç¨‹éƒ¨,ç‰¹ä¼‘,7\næå››,æ¥­å‹™éƒ¨,ç—…å‡,3\nç‹äº”,äººè³‡éƒ¨,äº‹å‡,2")
    csv_path = f.name

text, meta = DocumentParser.parse(csv_path, "csv")
test("CSV è§£æ: tables_detected = 1",    meta.get("tables_detected") == 1)
test("CSV è§£æ: åŒ…å« 'å¼µä¸‰'",           "å¼µä¸‰" in text)
test("CSV è§£æ: åŒ…å« 'å·¥ç¨‹éƒ¨'",          "å·¥ç¨‹éƒ¨" in text)
os.unlink(csv_path)

# 2.4 HTML è§£æ
with tempfile.NamedTemporaryFile(suffix=".html", mode="w", encoding="utf-8", delete=False) as f:
    f.write("""<html><head><title>å…¬å¸è¦å®š</title><style>body{}</style></head>
<body><h1>å·¥ä½œè¦å‰‡</h1><p>æœ¬è¦å‰‡é©ç”¨æ–¼å…¨é«”å“¡å·¥ã€‚</p>
<h2>å‡ºå‹¤ç®¡ç†</h2><p>ä¸Šåˆä¹é»ä¸Šç­ã€‚</p>
<table><tr><th>å‡åˆ¥</th><th>å¤©æ•¸</th></tr><tr><td>ç‰¹ä¼‘</td><td>7</td></tr></table>
<script>alert('test')</script></body></html>""")
    html_path = f.name

text, meta = DocumentParser.parse(html_path, "html")
test("HTML è§£æ: åŒ…å« 'å·¥ä½œè¦å‰‡'",        "å·¥ä½œè¦å‰‡" in text)
test("HTML è§£æ: ä¸åŒ…å« script å…§å®¹",     "alert" not in text)
test("HTML è§£æ: tables_detected >= 1",   meta.get("tables_detected", 0) >= 1)
os.unlink(html_path)

# 2.5 JSON è§£æ
with tempfile.NamedTemporaryFile(suffix=".json", mode="w", encoding="utf-8", delete=False) as f:
    json.dump({
        "company": "æ¸¬è©¦å…¬å¸",
        "policies": [
            {"name": "è«‹å‡è¾¦æ³•", "content": "ç—…å‡ä¸‰åæ—¥"},
            {"name": "åŠ ç­è¾¦æ³•", "content": "å¹³æ—¥åŠ ç­è²» 1.34 å€"}
        ]
    }, f, ensure_ascii=False)
    json_path = f.name

text, meta = DocumentParser.parse(json_path, "json")
test("JSON è§£æ: åŒ…å« 'æ¸¬è©¦å…¬å¸'",        "æ¸¬è©¦å…¬å¸" in text)
test("JSON è§£æ: åŒ…å« 'ç—…å‡ä¸‰åæ—¥'",       "ç—…å‡ä¸‰åæ—¥" in text)
os.unlink(json_path)

# 2.6 Excel è§£æ
if _HAS_OPENPYXL:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "è–ªè³‡è¡¨"
    ws.append(["è·ç´š", "åŸºæœ¬è–ªè³‡", "åŠ ç­è²»ç‡"])
    ws.append(["å·¥ç¨‹å¸«", "50000", "1.34"])
    ws.append(["è³‡æ·±å·¥ç¨‹å¸«", "65000", "1.34"])
    ws.append(["ä¸»ç®¡", "80000", "1.67"])
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        xlsx_path = f.name
    wb.save(xlsx_path)

    text, meta = DocumentParser.parse(xlsx_path, "xlsx")
    test("Excel è§£æ: tables_detected >= 1",  meta.get("tables_detected", 0) >= 1)
    test("Excel è§£æ: åŒ…å« 'å·¥ç¨‹å¸«'",         "å·¥ç¨‹å¸«" in text)
    test("Excel è§£æ: åŒ…å« '50000'",          "50000" in text)
    os.unlink(xlsx_path)
else:
    warn("openpyxl æœªå®‰è£ï¼Œè·³é Excel è§£ææ¸¬è©¦")

# 2.7 RTF è§£æ
if _HAS_RTF:
    with tempfile.NamedTemporaryFile(suffix=".rtf", mode="w", encoding="utf-8", delete=False) as f:
        f.write(r"{\rtf1\ansi\deff0{\fonttbl{\f0 Times New Roman;}}{\pard This is a test document about leave policy.\par}}")
        rtf_path = f.name

    text, meta = DocumentParser.parse(rtf_path, "rtf")
    test("RTF è§£æ: åŒ…å« 'leave policy'",  "leave policy" in text.lower())
    os.unlink(rtf_path)
else:
    warn("striprtf æœªå®‰è£ï¼Œè·³é RTF è§£ææ¸¬è©¦")

# 2.8 DOCX è§£æ
from docx import Document as DocxDoc
doc = DocxDoc()
doc.add_heading("å…¬å¸å·¥ä½œè¦å‰‡", level=1)
doc.add_paragraph("æœ¬è¦å‰‡é©ç”¨æ–¼æœ¬å…¬å¸å…¨é«”å“¡å·¥ã€‚")
doc.add_heading("ç¬¬ä¸€ç«  è–ªè³‡", level=2)
doc.add_paragraph("åŸºæœ¬è–ªè³‡ä¸å¾—ä½æ–¼æ³•å®šæœ€ä½å·¥è³‡ã€‚")
table = doc.add_table(rows=3, cols=2)
table.cell(0, 0).text = "é …ç›®"
table.cell(0, 1).text = "é‡‘é¡"
table.cell(1, 0).text = "åŸºæœ¬è–ªè³‡"
table.cell(1, 1).text = "27470"
table.cell(2, 0).text = "åŠ ç­è²»"
table.cell(2, 1).text = "ä¾æ³•è¨ˆç®—"
with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    docx_path = f.name
doc.save(docx_path)

text, meta = DocumentParser.parse(docx_path, "docx")
test("DOCX è§£æ: åŒ…å«æ¨™é¡Œ 'å·¥ä½œè¦å‰‡'",     "å·¥ä½œè¦å‰‡" in text)
test("DOCX è§£æ: tables_detected >= 1",    meta.get("tables_detected", 0) >= 1)
test("DOCX è§£æ: åŒ…å«è¡¨æ ¼å…§å®¹ '27470'",    "27470" in text)
test("DOCX è§£æ: åµæ¸¬åˆ°æ¨™é¡Œå±¤ç´š '#'",       "#" in text)
os.unlink(docx_path)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# 3. QualityReport æ¸¬è©¦
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

section("3. QualityReport â€” å“è³ªå ±å‘Šç³»çµ±")

report = QualityReport(format_detected="pdf", total_chars=5000, total_pages=10)
report.compute_quality()
test("å“è³ªåˆ†æ•¸ = 1.0 (ç„¡å•é¡Œ)", report.quality_score == 1.0)
test("å“è³ªç­‰ç´š = excellent",    report.quality_level == "excellent")

report2 = QualityReport(format_detected="pdf", total_chars=50)
report2.add_warning("ç¬¬ 3 é ç‚ºæƒæ")
report2.add_warning("è¡¨æ ¼æ ¼å¼å¯èƒ½éºå¤±")
report2.compute_quality()
test("æœ‰ 2 è­¦å‘Š: quality_score < 1.0", report2.quality_score < 1.0)
test("å“è³ªç­‰ç´šä¸ç‚º excellent",         report2.quality_level != "excellent")

report3 = QualityReport(format_detected="pdf", total_chars=10)
report3.add_error("OCR å¤±æ•—")
report3.add_error("ç„¡æ³•æå–æ–‡å­—")
report3.compute_quality()
test("æœ‰éŒ¯èª¤+ä½å­—æ•¸: quality_level = poor æˆ– failed", report3.quality_level in ("poor", "failed"))

d = report.to_dict()
test("to_dict() åŒ…å« quality_score",   "quality_score" in d)
test("to_dict() åŒ…å« format_detected",  "format_detected" in d)
test("to_dict() åŒ…å« warnings",         "warnings" in d)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# 4. TextChunker æ¸¬è©¦
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

section("4. TextChunker â€” æ™ºæ…§åˆ‡ç‰‡")

# 4.1 Token è¨ˆç®—
token_count = TextChunker.count_tokens("Hello world. ä½ å¥½ä¸–ç•Œã€‚")
test(f"Token è¨ˆç®—: 'Hello world. ä½ å¥½ä¸–ç•Œã€‚' = {token_count} tokens", token_count > 0)

if _HAS_TIKTOKEN:
    test("ä½¿ç”¨ tiktoken ç²¾ç¢ºè¨ˆç®—", True)
else:
    warn("tiktoken æœªå®‰è£ï¼Œä½¿ç”¨ä¼°ç®—æ¨¡å¼")

# 4.2 åŸºæœ¬åˆ‡ç‰‡
long_text = "\n\n".join([
    f"ç¬¬{i}æ¢ é€™æ˜¯ä¸€æ®µé—œæ–¼å“¡å·¥ç®¡ç†çš„æ¸¬è©¦æ–‡å­—ï¼ŒåŒ…å«å„ç¨®è¦å®šå’Œç´°ç¯€èªªæ˜ã€‚" * 5
    for i in range(1, 51)
])
chunks = TextChunker.split_by_tokens(long_text, chunk_size=200, chunk_overlap=30)
test(f"é•·æ–‡æœ¬åˆ‡ç‰‡: ç”¢ç”Ÿ {len(chunks)} å€‹ chunks > 0",  len(chunks) > 0)
test("æ¯å€‹ chunk ä¸ç‚ºç©º",  all(len(c.strip()) > 0 for c in chunks))

# 4.3 è¡¨æ ¼ä¿è­·
table_text = """# è–ªè³‡è¦å®š

åŸºæœ¬è–ªè³‡è¦å®šå¦‚ä¸‹ï¼š

[è¡¨æ ¼ 1]
è·ç´š | è–ªè³‡ | åŠ ç­è²»ç‡
å·¥ç¨‹å¸« | 50000 | 1.34
è³‡æ·±å·¥ç¨‹å¸« | 65000 | 1.34
ä¸»ç®¡ | 80000 | 1.67

ä»¥ä¸Šè¡¨æ ¼ç‚ºå…¬å¸è–ªè³‡æ¨™æº–ã€‚

# è«‹å‡è¦å®š

å“¡å·¥è«‹å‡éœ€æå‰ç”³è«‹ã€‚"""

table_chunks = TextChunker.split_by_tokens(table_text, chunk_size=500, chunk_overlap=50)
# æª¢æŸ¥è¡¨æ ¼æ˜¯å¦è¢«ä¿è­·
table_intact = any("[è¡¨æ ¼ 1]" in c and "ä¸»ç®¡" in c for c in table_chunks)
test("è¡¨æ ¼ä¿è­·: è¡¨æ ¼å…§å®¹åœ¨åŒä¸€ chunk ä¸­", table_intact)

# 4.4 æ¨™é¡Œé‚Šç•Œåˆ‡åˆ†
heading_text = """# ç¬¬ä¸€ç«  ç¸½å‰‡

æœ¬ç« èªªæ˜åŸºæœ¬è¦å®šã€‚

# ç¬¬äºŒç«  å‡ºå‹¤

ä¸Šåˆä¹é»ä¸Šç­ã€‚

# ç¬¬ä¸‰ç«  è«‹å‡

è«‹å‡éœ€æå‰ç”³è«‹ã€‚"""

heading_chunks = TextChunker.split_by_tokens(heading_text, chunk_size=500, chunk_overlap=50)
test(f"æ¨™é¡Œåˆ‡åˆ†: ç”¢ç”Ÿ {len(heading_chunks)} å€‹ chunks", len(heading_chunks) >= 1)

# 4.5 ç©ºæ–‡æœ¬
empty_chunks = TextChunker.split_by_tokens("", chunk_size=500)
test("ç©ºæ–‡æœ¬: è¿”å› []", empty_chunks == [])

# 4.6 å¾ˆçŸ­çš„æ–‡æœ¬ï¼ˆä½æ–¼æœ€ä½é–€æª»ï¼‰
short_chunks = TextChunker.split_by_tokens("çŸ­", chunk_size=500)
test("éçŸ­æ–‡æœ¬: è¢«éæ¿¾", short_chunks == [])


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# 5. KnowledgeBaseRetriever çµæ§‹æ¸¬è©¦
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

section("5. KnowledgeBaseRetriever â€” çµæ§‹é©—è­‰")

from app.services.kb_retrieval import KnowledgeBaseRetriever, _HAS_BM25

test("BM25 library å·²å®‰è£", _HAS_BM25)

# é©—è­‰é¡çš„æ–¹æ³•
test("æœ‰ search æ–¹æ³•",            hasattr(KnowledgeBaseRetriever, "search"))
test("æœ‰ batch_search æ–¹æ³•",      hasattr(KnowledgeBaseRetriever, "batch_search"))
test("æœ‰ get_stats æ–¹æ³•",         hasattr(KnowledgeBaseRetriever, "get_stats"))
test("æœ‰ _semantic_search æ–¹æ³•",  hasattr(KnowledgeBaseRetriever, "_semantic_search"))
test("æœ‰ _keyword_search æ–¹æ³•",   hasattr(KnowledgeBaseRetriever, "_keyword_search"))
test("æœ‰ _hybrid_search æ–¹æ³•",    hasattr(KnowledgeBaseRetriever, "_hybrid_search"))
test("æœ‰ _rerank æ–¹æ³•",           hasattr(KnowledgeBaseRetriever, "_rerank"))
test("æœ‰ invalidate_cache æ–¹æ³•",  hasattr(KnowledgeBaseRetriever, "invalidate_cache"))

# é©—è­‰ search æ–¹æ³•ç°½å
import inspect
sig = inspect.signature(KnowledgeBaseRetriever.search)
params = list(sig.parameters.keys())
test("search() æœ‰ mode åƒæ•¸",      "mode" in params)
test("search() æœ‰ min_score åƒæ•¸",  "min_score" in params)
test("search() æœ‰ rerank åƒæ•¸",     "rerank" in params)
test("search() æœ‰ use_cache åƒæ•¸",  "use_cache" in params)

# BM25 åˆ†è©æ¸¬è©¦
tokenize = KnowledgeBaseRetriever._tokenize
tokens = tokenize("å“¡å·¥è«‹å‡è¾¦æ³• employee leave policy")
test(f"ä¸­è‹±æ··åˆåˆ†è©: ç”¢ç”Ÿ {len(tokens)} å€‹ tokens > 5", len(tokens) > 5)
test("åŒ…å«ä¸­æ–‡å­—å…ƒ",  "å“¡" in tokens)
test("åŒ…å«è‹±æ–‡è©",    "employee" in tokens)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# 6. è¨­å®šæª¢æŸ¥
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

section("6. è¨­å®š & Schema å®Œæ•´æ€§")

from app.config import settings

test("settings æœ‰ RETRIEVAL_MODE",       hasattr(settings, "RETRIEVAL_MODE"))
test("settings æœ‰ RETRIEVAL_MIN_SCORE",  hasattr(settings, "RETRIEVAL_MIN_SCORE"))
test("settings æœ‰ RETRIEVAL_RERANK",     hasattr(settings, "RETRIEVAL_RERANK"))
test("settings æœ‰ RETRIEVAL_CACHE_TTL",  hasattr(settings, "RETRIEVAL_CACHE_TTL"))
test("settings æœ‰ RETRIEVAL_TOP_K",      hasattr(settings, "RETRIEVAL_TOP_K"))
test("RETRIEVAL_MODE = 'hybrid'",        settings.RETRIEVAL_MODE == "hybrid")

from app.schemas.document import DocumentUpdate, Document

du = DocumentUpdate(quality_report={"quality_score": 0.9, "warnings": []})
test("DocumentUpdate æ”¯æ´ quality_report", du.quality_report is not None)

from app.models.document import DocumentChunk
test("DocumentChunk æœ‰ vector_id æ¬„ä½", hasattr(DocumentChunk, "vector_id"))

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# 7. æ•ˆèƒ½æ¸¬è©¦
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

section("7. æ•ˆèƒ½åŸºæº–æ¸¬è©¦")

# å¤§æ–‡ä»¶åˆ‡ç‰‡æ•ˆèƒ½
big_text = ("é€™æ˜¯ä¸€æ®µå¾ˆé•·çš„æ¸¬è©¦æ–‡å­—ï¼Œç”¨ä¾†æ¨¡æ“¬å¤§å‹ä¼æ¥­æ–‡ä»¶çš„å…§å®¹ã€‚" * 100 + "\n\n") * 50
start_time = time.time()
big_chunks = TextChunker.split_by_tokens(big_text, chunk_size=1000, chunk_overlap=150)
elapsed = (time.time() - start_time) * 1000
test(f"å¤§æ–‡ä»¶åˆ‡ç‰‡ ({len(big_text)} å­—): {elapsed:.0f}ms < 5000ms", elapsed < 5000)
test(f"å¤§æ–‡ä»¶åˆ‡ç‰‡: ç”¢ç”Ÿ {len(big_chunks)} å€‹ chunks", len(big_chunks) > 0)

# TXT è§£ææ•ˆèƒ½
with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False) as f:
    f.write(big_text)
    big_txt_path = f.name
start_time = time.time()
text, meta = DocumentParser.parse(big_txt_path, "txt")
parse_elapsed = (time.time() - start_time) * 1000
test(f"å¤§ TXT è§£æ: {parse_elapsed:.0f}ms < 3000ms", parse_elapsed < 3000)
os.unlink(big_txt_path)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ç¸½çµ
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

section("æ¸¬è©¦ç¸½çµ")

total = passed + failed
print(f"\n  é€šé: {GREEN}{passed}{RESET}")
print(f"  å¤±æ•—: {RED}{failed}{RESET}")
print(f"  è­¦å‘Š: {YELLOW}{warnings}{RESET}")
print(f"  é€šéç‡: {GREEN if failed == 0 else YELLOW}{passed}/{total} ({passed/total*100:.1f}%){RESET}")
print()

if failed == 0:
    print(f"  {GREEN}{BOLD}â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•—{RESET}")
    print(f"  {GREEN}{BOLD}â•‘            âœ“ ALL TESTS PASSED                    â•‘{RESET}")
    print(f"  {GREEN}{BOLD}â•šâ•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•{RESET}")
else:
    print(f"  {RED}{BOLD}â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•—{RESET}")
    print(f"  {RED}{BOLD}â•‘         âœ— {failed} TEST(S) FAILED                     â•‘{RESET}")
    print(f"  {RED}{BOLD}â•šâ•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•{RESET}")

print(f"""
{CYAN}æ–°å¢èƒ½åŠ›æ‘˜è¦ï¼š{RESET}
  ğŸ“„ æ–‡ä»¶æ ¼å¼: {len(SUPPORTED_FORMATS)} ç¨®ï¼ˆPDF/DOCX/DOC/TXT/Excel/CSV/HTML/MD/RTF/JSON/åœ–ç‰‡ï¼‰
  ğŸ”¢ Token è¨ˆç®—: {'tiktoken ç²¾ç¢ºè¨ˆç®—' if _HAS_TIKTOKEN else 'ä¼°ç®—æ¨¡å¼'}
  ğŸ“Š PDF è¡¨æ ¼: {'pdfplumber å•Ÿç”¨' if _HAS_PDFPLUMBER else 'æœªå•Ÿç”¨'}
  ğŸ“‹ Excel: {'openpyxl å•Ÿç”¨' if _HAS_OPENPYXL else 'æœªå•Ÿç”¨'}
  ğŸ” OCR: {'pytesseract å•Ÿç”¨' if _HAS_OCR else 'æœªå®‰è£ï¼ˆéœ€ç³»çµ±å®‰è£ tesseractï¼‰'}
  ğŸ”  ç·¨ç¢¼åµæ¸¬: {'chardet å•Ÿç”¨' if _HAS_CHARDET else 'æœªå•Ÿç”¨'}
  ğŸ“‘ RTF: {'striprtf å•Ÿç”¨' if _HAS_RTF else 'æœªå•Ÿç”¨'}
  ğŸ” BM25 é—œéµå­—æª¢ç´¢: {'rank-bm25 å•Ÿç”¨' if _HAS_BM25 else 'æœªå•Ÿç”¨'}
  ğŸ”„ æ··åˆæª¢ç´¢ (RRF): å·²å¯¦ç¾
  âš¡ é‡æ’åº (Voyage Rerank): å·²å¯¦ç¾
  ğŸ—„ï¸ Redis æŸ¥è©¢å¿«å–: å·²å¯¦ç¾
""")

if __name__ == "__main__":
    sys.exit(0 if failed == 0 else 1)
